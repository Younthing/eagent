"""Docx report renderer for ROB2."""

from __future__ import annotations

from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from reporting.context import DOMAIN_NAMES, RISK_LABELS, build_report_context
from schemas.responses import Rob2RunResult

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument


def generate_docx_report(result: Rob2RunResult, output_path, pdf_name: str = "Unknown") -> None:
    """Generate Word (docx) report."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    heading = doc.add_heading("ROB2 风险偏倚评估报告", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    context = build_report_context(result, pdf_name=pdf_name)
    doc.add_paragraph(f"PDF 文件: {pdf_name}")
    doc.add_paragraph(f"生成时间: {context['generated_at']}")

    doc.add_heading("文献元信息", level=1)
    _add_metadata_section(doc, context.get("document_metadata"))

    doc.add_heading("总体评估", level=1)
    _add_overall_section(doc, result)

    doc.add_heading("领域评估详情", level=1)

    for domain in result.result.domains:
        _add_domain_section(doc, domain)
        doc.add_page_break()

    doc.save(str(output_path))


def _add_overall_section(doc: "DocxDocument", result: Rob2RunResult) -> None:
    p = doc.add_paragraph()
    run = p.add_run(
        f"总体风险: {RISK_LABELS.get(result.result.overall.risk, result.result.overall.risk)}"
    )
    run.bold = True
    p.add_run(f"\n{result.result.overall.rationale}")


def _add_metadata_section(doc: "DocxDocument", metadata: dict | None) -> None:
    if not metadata:
        doc.add_paragraph("未提取到文献元信息。")
        return

    doc.add_paragraph(f"标题: {metadata.get('title') or ''}")
    doc.add_paragraph(f"作者: {_join_values(metadata.get('authors'))}")
    doc.add_paragraph(f"年份: {metadata.get('year') or ''}")
    doc.add_paragraph(f"机构: {_join_values(metadata.get('affiliations'))}")
    doc.add_paragraph(f"基金: {_join_values(metadata.get('funders'))}")

    extraction = metadata.get("extraction")
    if isinstance(extraction, dict):
        doc.add_paragraph(
            "提取信息: "
            f"method={extraction.get('method') or ''}, "
            f"model={extraction.get('model_id') or ''}, "
            f"provider={extraction.get('provider') or ''}, "
            f"confidence={extraction.get('confidence') if extraction.get('confidence') is not None else ''}, "
            f"error={extraction.get('error') or ''}"
        )

    sources = metadata.get("sources")
    if isinstance(sources, list) and sources:
        p = doc.add_paragraph()
        p.add_run("元信息来源:").bold = True
        for source in sources:
            if not isinstance(source, dict):
                continue
            paragraph_id = source.get("paragraph_id") or ""
            quote = source.get("quote") or ""
            doc.add_paragraph(f"- {paragraph_id}: {quote}", style="List Bullet")


def _add_domain_section(doc: "DocxDocument", domain) -> None:
    domain_name = DOMAIN_NAMES.get(domain.domain, domain.domain)
    risk_label = RISK_LABELS.get(domain.risk, domain.risk)

    doc.add_heading(f"{domain_name} - {risk_label}", level=2)
    doc.add_paragraph(f"风险理由: {domain.risk_rationale}")

    if getattr(domain, "rule_trace", None):
        p = doc.add_paragraph()
        p.add_run("规则路径:").bold = True
        for trace in domain.rule_trace:
            doc.add_paragraph(f"- {trace}", style="List Bullet")

    for answer in domain.answers:
        p = doc.add_paragraph()
        p.add_run(f"{answer.question_id}: {answer.text or 'Question text unavailable'}").bold = True

        p = doc.add_paragraph()
        p.add_run("回答: ").bold = True
        p.add_run(f"{answer.answer}")

        p = doc.add_paragraph()
        p.add_run("理由: ").italic = True
        p.add_run(f"{answer.rationale}")

        if answer.evidence_refs:
            p = doc.add_paragraph()
            p.add_run("证据引用:").bold = True
            for ref in answer.evidence_refs:
                doc.add_paragraph(
                    f"- {ref.quote} (Page {ref.page or '?'})",
                    style="List Bullet",
                )


def _join_values(values: object) -> str:
    if not isinstance(values, list):
        return ""
    parts = []
    for item in values:
        if item is None:
            continue
        text = str(item).strip()
        if text:
            parts.append(text)
    return "; ".join(parts)


__all__ = ["generate_docx_report"]
