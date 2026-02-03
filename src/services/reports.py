"""Report generation services for ROB2."""

from __future__ import annotations

import datetime
from pathlib import Path

from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from schemas.responses import Rob2RunResult

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False


RISK_LABELS = {
    "low": "低风险",
    "some_concerns": "有些疑虑",
    "high": "高风险",
    "not_applicable": "不适用"
}

DOMAIN_NAMES = {
    "D1": "D1: 随机化过程产生的偏差",
    "D2": "D2: 偏离既定干预措施产生的偏差",
    "D3": "D3: 缺失结果数据产生的偏差",
    "D4": "D4: 结果测量产生的偏差",
    "D5": "D5: 结果报告选择产生的偏差",
}

def generate_html_report(result: Rob2RunResult, output_path: Path, pdf_name: str = "Unknown") -> None:
    """Generate interactive HTML report."""
    template_dir = Path(__file__).parent.parent / "templates"
    env = Environment(loader=FileSystemLoader(str(template_dir)))
    template = env.get_template("report.html")

    html_content = template.render(
        data=result.result,
        pdf_name=pdf_name,
        generated_at=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        risk_labels=RISK_LABELS,
        domain_names=DOMAIN_NAMES
    )

    output_path.write_text(html_content, encoding="utf-8")


def generate_pdf_report(result: Rob2RunResult, output_path: Path, pdf_name: str = "Unknown") -> None:
    """Generate PDF report from HTML."""
    if not WEASYPRINT_AVAILABLE:
        print("Warning: WeasyPrint not installed, skipping PDF generation.")
        return

    template_dir = Path(__file__).parent.parent / "templates"
    env = Environment(loader=FileSystemLoader(str(template_dir)))
    template = env.get_template("report.html")

    html_content = template.render(
        data=result.result,
        pdf_name=pdf_name,
        generated_at=datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        risk_labels=RISK_LABELS,
        domain_names=DOMAIN_NAMES
    )

    # We need to make sure all sections are expanded for PDF
    # The CSS @media print should handle it, but weasyprint is a browser engine so it respects CSS.
    # However, we might want to inject specific CSS to force display block if media query fails

    css = CSS(string="""
        .domain-content { display: block !important; }
        .container { box-shadow: none; padding: 0; }
        body { background-color: white; }
    """)

    HTML(string=html_content, base_url=str(template_dir)).write_pdf(output_path, stylesheets=[css])


def generate_docx_report(result: Rob2RunResult, output_path: Path, pdf_name: str = "Unknown") -> None:
    """Generate Word (docx) report."""
    doc = Document()

    # Set default font to something that supports Chinese well if needed,
    # but python-docx defaults are usually okay for basic text.
    # To be safe for Chinese, we often need to set eastAsia font.
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Title
    heading = doc.add_heading('ROB2 风险偏倚评估报告', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph(f"PDF 文件: {pdf_name}")
    doc.add_paragraph(f"生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('总体评估', level=1)

    # Overall Risk
    p = doc.add_paragraph()
    run = p.add_run(f"总体风险: {RISK_LABELS.get(result.result.overall.risk, result.result.overall.risk)}")
    run.bold = True
    p.add_run(f"\n{result.result.overall.rationale}")

    doc.add_heading('领域评估详情', level=1)

    for domain in result.result.domains:
        domain_name = DOMAIN_NAMES.get(domain.domain, domain.domain)
        risk_label = RISK_LABELS.get(domain.risk, domain.risk)

        doc.add_heading(f"{domain_name} - {risk_label}", level=2)

        doc.add_paragraph(f"风险理由: {domain.risk_rationale}")

        for answer in domain.answers:
            p = doc.add_paragraph()
            p.add_run(f"{answer.question_id}: {answer.text or 'Question text unavailable'}").bold = True

            p = doc.add_paragraph()
            p.add_run("回答: ").bold = True
            p.add_run(f"{answer.answer}")

            p = doc.add_paragraph()
            p.add_run("理由: ").italic = True
            p.add_run(f"{answer.rationale}")

            if answer.evidence_refs:
                p = doc.add_paragraph()
                p.add_run("证据引用:").bold = True
                for ref in answer.evidence_refs:
                    doc.add_paragraph(f"- {ref.quote} (Page {ref.page or '?'})", style='List Bullet')

        doc.add_page_break()

    doc.save(str(output_path))
